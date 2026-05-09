"""
scripts/make_run_doc.py
Generate LevelUp_RunCommands.docx using python-docx.
Run from project root: python scripts/make_run_doc.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUTPUT = os.path.join(
    os.path.dirname(__file__), "..", "..", "LevelUp_RunCommands.docx"
)
OUTPUT = os.path.normpath(OUTPUT)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def shade_paragraph(para, hex_color="F0F0F0"):
    """Add paragraph-level shading (simulates code block background)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Cambria"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_code_block(doc, lines):
    """Add a shaded code block. lines can be a string or list of strings."""
    if isinstance(lines, str):
        lines = lines.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        shade_paragraph(p, "F2F2F2")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths_in=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Calibri"
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr_cells[i], "1F3964")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_borders(hdr_cells[i], "1F3964")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "FFFFFF" if r_idx % 2 == 0 else "EBF3FB"
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            row_cells[c_idx].paragraphs[0].runs[0].font.name = "Calibri"
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(row_cells[c_idx], bg)
            set_cell_borders(row_cells[c_idx], "CCCCCC")

    # Column widths
    if col_widths_in:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths_in[i])

    doc.add_paragraph()  # spacer


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins (1 inch)
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

# Default body font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── TITLE ─────────────────────────────────────────────────────────────────────

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("⚔️  LevelUp AI — Project Run Commands")
title_run.font.name = "Cambria"
title_run.font.size = Pt(22)
title_run.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Applied AI Exam — SRH Stuttgart")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(13)
sub_run.italic = True
sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()  # spacer

# ── SECTION 1: Prerequisites ──────────────────────────────────────────────────

add_heading(doc, "1. Prerequisites")
add_body(doc, "Ensure the following tools are installed before running the project:")
add_bullet(doc, "Python 3.10+ — https://python.org/downloads")
add_bullet(doc, "Git — https://git-scm.com")
add_bullet(doc, "NVIDIA GPU with CUDA 12.8 support (RTX 4060 or better recommended)")
add_bullet(doc, "8 GB+ VRAM for model inference")
add_bullet(doc, "A HuggingFace account with token (to download LLaMA 3.2)")
add_bullet(doc, "ElevenLabs API key (optional — enables per-build AI voices)")
doc.add_paragraph()

# ── SECTION 2: First-Time Setup ───────────────────────────────────────────────

add_heading(doc, "2. First-Time Setup")
add_body(doc, "Run these commands once after cloning the repository.")

add_heading(doc, "2.1  Clone the Repository", level=2)
add_code_block(doc, [
    "git clone https://github.com/veeravenkatsaikondaiahpalpu-hue/levelup-ai.git",
    "cd levelup-ai",
])

add_heading(doc, "2.2  Create & Activate Virtual Environment", level=2)
add_body(doc, "Windows:")
add_code_block(doc, [
    "python -m venv venv",
    r"venv\Scripts\activate",
])
add_body(doc, "Mac / Linux:")
add_code_block(doc, [
    "python -m venv venv",
    "source venv/bin/activate",
])

add_heading(doc, "2.3  Install PyTorch with CUDA 12.8 (install FIRST)", level=2)
add_code_block(doc, [
    "pip install torch==2.11.0+cu128 --index-url https://download.pytorch.org/whl/cu128",
])

add_heading(doc, "2.4  Install Remaining Dependencies", level=2)
add_code_block(doc, [
    "pip install -r requirements.txt",
])
doc.add_paragraph()

# ── SECTION 3: Environment Configuration ─────────────────────────────────────

add_heading(doc, "3. Environment Configuration")
add_body(doc, 'Create a file named  .env  in the project root with the following contents:')
add_code_block(doc, [
    "HF_TOKEN=your_huggingface_token_here",
    "LEVELUP_LOAD_MODEL=true",
    "ELEVENLABS_API_KEY=your_elevenlabs_key_here   # optional",
    "WHISPER_MODEL=base                            # optional: base / small / medium",
])
add_body(doc, "Variable reference:")
add_table(doc,
    ["Variable", "Required", "Description", "Where to get it"],
    [
        ["HF_TOKEN", "Yes", "HuggingFace access token to download LLaMA 3.2", "huggingface.co/settings/tokens"],
        ["LEVELUP_LOAD_MODEL", "Yes", "Set to true to load the fine-tuned model on startup", "Set to false for frontend-only dev"],
        ["ELEVENLABS_API_KEY", "No", "Enables per-build AI voices via ElevenLabs TTS", "elevenlabs.io"],
        ["WHISPER_MODEL", "No", "Whisper model size for local STT fallback", "base / small / medium"],
    ],
    col_widths_in=[1.5, 0.9, 2.8, 1.8]
)

# ── SECTION 4: Running the Project ───────────────────────────────────────────

add_heading(doc, "4. Running the Project")
add_body(doc, "Open two separate terminal windows / command prompts.")

add_heading(doc, "Terminal 1 — Backend API  (port 8000)", level=2)
add_code_block(doc, [
    "# Make sure venv is activated first",
    r"venv\Scripts\activate          # Windows",
    "# source venv/bin/activate      # Mac/Linux",
    "",
    "uvicorn api.main:app --port 8000 --reload",
])
add_body(doc, 'Wait for the message:  [LevelUp API] Model ready.  — the chatbot is now live.', bold=False)

add_heading(doc, "Terminal 2 — Frontend  (port 8080)", level=2)
add_code_block(doc, [
    "cd frontend",
    "python -m http.server 8080",
])
add_body(doc, "Then open your browser and navigate to:")
add_code_block(doc, ["http://localhost:8080/app.html"])
doc.add_paragraph()

# ── SECTION 5: API Endpoints ──────────────────────────────────────────────────

add_heading(doc, "5. API Endpoints")
add_body(doc, "All endpoints are served on  http://localhost:8000  when the backend is running.")
add_table(doc,
    ["Endpoint", "Method", "Description"],
    [
        ["/api/chat",             "POST", "Fine-tuned LLM chatbot — send message, get build-persona reply"],
        ["/api/activity/check",  "POST", "Anomaly detection on user activity patterns (XP fraud check)"],
        ["/api/sentiment/analyze","POST", "Real-time sentiment classifier from chat input"],
        ["/api/tts",             "POST", "ElevenLabs TTS — build-specific voice, returns MP3 audio"],
        ["/api/stt",             "POST", "OpenAI Whisper STT — audio file → transcribed text"],
        ["/api/voices",          "GET",  "List all per-build ElevenLabs voice profiles"],
        ["/health",              "GET",  "Server health check + model load status"],
        ["/docs",                "GET",  "Interactive Swagger / OpenAPI documentation"],
    ],
    col_widths_in=[2.0, 0.8, 4.2]
)

# ── SECTION 6: Voice Features ─────────────────────────────────────────────────

add_heading(doc, "6. Voice Features")

add_heading(doc, "Speech-to-Text (STT) — Mic Input", level=2)
add_bullet(doc, "Click the 🎤 mic button in the AI Coach chat panel")
add_bullet(doc, "Primary: Browser Web Speech API (zero setup, real-time)")
add_bullet(doc, "Fallback: OpenAI Whisper via /api/stt (if browser STT unavailable)")
add_bullet(doc, "Transcribed speech is auto-filled and sent as a message")
doc.add_paragraph()

add_heading(doc, "Text-to-Speech (TTS) — AI Voice Output", level=2)
add_bullet(doc, "Click the 🔊 speaker icon in the chat header to toggle voice responses")
add_bullet(doc, "Primary: ElevenLabs /api/tts when ELEVENLABS_API_KEY is set")
add_bullet(doc, "Fallback: browser speechSynthesis API (no key needed)")
doc.add_paragraph()

add_body(doc, "Per-build voice profiles:")
add_table(doc,
    ["Build", "Voice Style"],
    [
        ["TITAN",   "Deep, intense, punchy — low stability, high drama"],
        ["ORACLE",  "Clear, measured, analytical — high stability, minimal style"],
        ["PHANTOM", "Fluid, energetic, dynamic — medium expressiveness"],
        ["SAGE",    "Ultra-calm, slow, grounded — highest stability, near-zero style"],
        ["MUSE",    "Warm, expressive, creative — lowest stability, most varied"],
        ["EMPIRE",  "Authoritative, commanding — high stability, professional"],
        ["GG",      "Hype, chaotic, gamer energy — max style, maximum variation"],
    ],
    col_widths_in=[1.2, 5.8]
)

# ── SECTION 7: Model Training ─────────────────────────────────────────────────

add_heading(doc, "7. Model Training (Optional)")
add_body(doc, "Only required if you want to retrain the QLoRA fine-tuned LLaMA 3.2 3B model.")

add_heading(doc, "Full training  (~8–12 hours on RTX 4060)", level=2)
add_code_block(doc, ["python -m chatbot.fine_tuning.train"])

add_heading(doc, "Train a single build", level=2)
add_code_block(doc, ["python -m chatbot.fine_tuning.train --build TITAN"])

add_heading(doc, "Quick smoke test  (~5 minutes)", level=2)
add_code_block(doc, ["python -m chatbot.fine_tuning.train --smoke_test"])

add_heading(doc, "Training Configuration", level=2)
add_table(doc,
    ["Parameter", "Value"],
    [
        ["Base model",    "unsloth/Llama-3.2-3B-Instruct"],
        ["Quantisation",  "QLoRA 4-bit NF4"],
        ["LoRA rank",     "r=16, α=32"],
        ["Target modules","All attention + MLP projections"],
        ["Trainer",       "TRL SFTTrainer with sequence packing"],
        ["Dataset size",  "76,616 samples across 7 builds"],
        ["Hardware",      "NVIDIA RTX 4060 Laptop GPU (8 GB VRAM)"],
    ],
    col_widths_in=[2.0, 5.0]
)

# ── SECTION 8: Troubleshooting ────────────────────────────────────────────────

add_heading(doc, "8. Troubleshooting")
add_table(doc,
    ["Problem", "Likely Cause", "Fix"],
    [
        ["Model not loading / 503 on /api/chat",
         "LEVELUP_LOAD_MODEL not set or adapter missing",
         "Set LEVELUP_LOAD_MODEL=true in .env and ensure models/final/ exists"],
        ["CUDA out of memory",
         "GPU VRAM < 8 GB or another process using GPU",
         "Close other GPU apps; use --load-in-4bit (default) or reduce batch"],
        ["TTS returns 503",
         "ELEVENLABS_API_KEY not set",
         "Add key to .env or toggle voice off — browser TTS fallback is automatic"],
        ["STT returns 503",
         "openai-whisper not installed",
         "pip install openai-whisper (browser Web Speech API still works)"],
        ["Frontend shows blank / CORS error",
         "Backend not running or wrong port",
         "Start uvicorn on port 8000 first, then serve frontend on 8080"],
        ["XP not updating after log",
         "Daily XP cap already reached (600 XP base)",
         "Check Daily Cap in the Log Activity modal; resets at midnight"],
        ["pip install torch fails",
         "Wrong CUDA version URL",
         "Confirm GPU supports CUDA 12.8; see pytorch.org/get-started for your version"],
    ],
    col_widths_in=[1.8, 1.9, 3.3]
)

# ── FOOTER note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_run = note.add_run(
    "Swagger API docs available at  http://localhost:8000/docs  when the backend is running."
)
note_run.font.name = "Calibri"
note_run.font.size = Pt(10)
note_run.italic = True
note_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"[OK] Saved: {OUTPUT}")
